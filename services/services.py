from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.exc import IntegrityError
from sqlalchemy.future import select
from deep_translator import GoogleTranslator
from uuid import UUID
from datetime import datetime
from fastapi import HTTPException
from db.models import UserModel, TextModel, ArticleModel, AIModel, FileModel
from utils.secrets import password_manager
from schema._output import RegisterOutput
from utils.jwt import JWTHandler
from utils.storage import save_file
from uuid import uuid4
from docx import Document
from fpdf import FPDF
from io import BytesIO
import numpy as np
import faiss
import bcrypt
import torch
import torch.nn as nn
from transformers import (
    AutoTokenizer,
    AutoModelForCausalLM,
    BitsAndBytesConfig,
    DebertaV2ForSequenceClassification,
    DebertaV2Tokenizer,
)
from peft import LoraConfig, get_peft_model, prepare_model_for_kbit_training
from sentence_transformers import SentenceTransformer
from datasets import load_dataset
import nltk
from nltk.tokenize import sent_tokenize
from rank_bm25 import BM25Okapi
from trl import SFTTrainer
import matplotlib.pyplot as plt
from pylatexenc.latexencode import latexencode
import requests
from evaluate import load
import difflib
import logging
from torch.nn.parallel import DistributedDataParallel as DDP
import os

nltk.download("punkt", quiet=True)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def hash_password(password: str) -> str:
    salt = bcrypt.gensalt()
    hashed_password = bcrypt.hashpw(password.encode("utf-8"), salt)
    return hashed_password.decode("utf-8")


def verify_password(stored_password: str, provided_password: str) -> bool:
    return bcrypt.checkpw(
        provided_password.encode("utf-8"), stored_password.encode("utf-8")
    )


class UserOperation:
    def __init__(self, db_session: AsyncSession) -> None:
        self.db_session = db_session

    async def register_user(
        self,
        username: str,
        phone_number: str,
        degree: str,
        password: str,
        confirm_password: str,
    ) -> RegisterOutput:
        if password != confirm_password:
            raise HTTPException(
                status_code=400,
                detail="Passwords do not match.",
            )

        async with self.db_session as session:
            try:
                existing_user = await session.execute(
                    select(UserModel).where(
                        (UserModel.username == username)
                        | (UserModel.phone_number == phone_number)
                    )
                )
                if existing_user.scalars().first():
                    raise HTTPException(
                        status_code=400,
                        detail="User with this username or phone number already exists.",
                    )

                new_user = UserModel(
                    username=username,
                    phone_number=phone_number,
                    password=password_manager.hash(password),
                    degree=degree,
                    recovery_code=None,
                    email=None,
                    age=None,
                    family=None,
                    profile_image=None,
                )
                session.add(new_user)
                await session.commit()
                return RegisterOutput(username=new_user.username, id=new_user.id)
            except IntegrityError:
                await session.rollback()
                raise HTTPException(
                    status_code=400,
                    detail="User with this username or phone number already exists.",
                )

    async def login_user(self, username: str, password: str) -> dict:
        async with self.db_session as session:
            user = await session.execute(
                select(UserModel).where(UserModel.username == username)
            )
            user = user.scalars().first()

            if user is None or not password_manager.verify(password, user.password):
                raise HTTPException(
                    status_code=400, detail="Invalid username or password."
                )

            token = JWTHandler.generate(username)
            return {"token": token, "user_id": user.id}

    async def get_user_by_username(self, username: str):
        stmt = select(UserModel).where(UserModel.username == username)
        result = await self.db_session.execute(stmt)
        user = result.scalar_one_or_none()

        if not user:
            raise HTTPException(status_code=404, detail="User not found.")

        return user

    async def get_user_profile(self, user_identifier: str) -> dict:
        async with self.db_session as session:
            query = select(UserModel).where(
                (UserModel.id == user_identifier)
                | (UserModel.username == user_identifier)
            )
            user = await session.execute(query)
            user = user.scalars().first()

            if not user:
                raise HTTPException(status_code=404, detail="User not found")

            return {
                "username": user.username,
                "name": user.name,
                "family": user.family,
                "profile_image": user.profile_image,
            }

    async def delete_user(self, username: str) -> None:
        async with self.db_session as session:
            user = await session.execute(
                select(UserModel).where(UserModel.username == username)
            )
            user = user.scalars().first()

            if user is None:
                raise HTTPException(status_code=404, detail="User not found.")

            await session.delete(user)
            await session.commit()

    async def update_user(
        self,
        username: str,
        family: str = None,
        age: int = None,
        phone_number: str = None,
        degree: str = None,
        email: str = None,
    ) -> UserModel:
        async with self.db_session as session:
            user = await session.execute(
                select(UserModel).where(UserModel.username == username)
            )
            user = user.scalars().first()

            if user is None:
                raise HTTPException(status_code=404, detail="User not found.")

            if family is not None:
                user.family = family
            if age is not None:
                user.age = age
            if phone_number is not None:
                user.phone_number = phone_number
            if degree is not None:
                user.degree = degree
            if email is not None:
                user.email = email

            await session.commit()
            return user

    async def reset_password_by_username(self, username: str, new_password: str):
        user = await self.get_user_by_username(username)
        if not user:
            raise HTTPException(status_code=404, detail="User not found.")

        hashed_password = hash_password(new_password)
        user.password = hashed_password
        await self.db_session.commit()
        return {"message": "Password updated successfully."}

    async def generate_recovery_code(self, username: str) -> str:
        stmt = select(UserModel).where(UserModel.username == username)
        result = await self.db_session.execute(stmt)
        user = result.scalar_one_or_none()

        if not user:
            raise HTTPException(status_code=404, detail="User not found.")

        recovery_code = str(uuid4()).split("-")[0]
        user.recovery_code = recovery_code
        await self.db_session.commit()
        return recovery_code

    async def verify_recovery_code_and_reset_password(
        self,
        username: str,
        recovery_code: str,
        new_password: str,
        user_operation_service,
    ):
        stmt = select(UserModel).where(UserModel.username == username)
        result = await self.db_session.execute(stmt)
        user = result.scalar_one_or_none()

        if not user or user.recovery_code != recovery_code:
            raise HTTPException(status_code=400, detail="Invalid recovery code.")

        await user_operation_service.reset_password_by_username(username, new_password)
        user.recovery_code = None
        await self.db_session.commit()
        return {"detail": "Password reset successfully."}


class TextOperation:
    def __init__(self, db_session: AsyncSession) -> None:
        self.db_session = db_session

    async def create_text(
        self, user_id: UUID, text_content: str, creation_date: datetime = None
    ) -> dict:
        async with self.db_session as session:
            user = await session.execute(
                select(UserModel).where(UserModel.id == user_id)
            )
            user = user.scalars().first()
            if user is None:
                raise HTTPException(status_code=404, detail="User not found.")

            if creation_date is None:
                creation_date = datetime.utcnow()

            new_text = TextModel(
                user_id=user.id, text=text_content, creation_date=creation_date
            )
            session.add(new_text)
            await session.commit()
            return {
                "success": "Text created successfully.",
                "text_id": str(new_text.id),
            }

    async def delete_text(self, text_id: UUID) -> None:
        async with self.db_session as session:
            text = await session.execute(
                select(TextModel).where(TextModel.id == text_id)
            )
            text = text.scalars().first()
            if text is None:
                raise HTTPException(status_code=404, detail="Text not found.")

            await session.delete(text)
            await session.commit()

    async def update_text(self, text_id: UUID, new_text: str) -> dict:
        async with self.db_session as session:
            text = await session.execute(
                select(TextModel).where(TextModel.id == text_id)
            )
            text = text.scalars().first()
            if text is None:
                raise HTTPException(status_code=404, detail="Text not found.")

            text.text = new_text
            await session.commit()
            return {"success": "Text updated successfully."}

    async def list_user_texts(self, user_id: UUID) -> list:
        async with self.db_session as session:
            result = await session.execute(
                select(TextModel).where(TextModel.user_id == user_id)
            )
            texts = result.scalars().all()

            if not texts:
                raise HTTPException(
                    status_code=404, detail="No texts found for the user."
                )

            result_list = []
            for text in texts:
                result_list.append(
                    {
                        "id": str(text.id),
                        "user_id": str(text.user_id),
                        "content": text.text,
                        "creation_date": text.creation_date,
                    }
                )
            return result_list


class CoherenceModel(nn.Module):
    def __init__(self, input_dim=768):
        super(CoherenceModel, self).__init__()
        self.fc = nn.Linear(input_dim, 1)
        self.sigmoid = nn.Sigmoid()

    def forward(self, x):
        return self.sigmoid(self.fc(x))


class KnowledgeBase:
    def __init__(self, db_session: AsyncSession) -> None:
        self.db_session = db_session
        self.embedder = SentenceTransformer("all-mpnet-base-v2")
        self.index = None
        self.id_map = []
        self.dimension = 768
        self.memory = []
        self.bm25 = None

    async def build_index(self, use_fallback_dataset=True):
        async with self.db_session as session:
            result = await session.execute(select(ArticleModel))
            articles = result.scalars().all()

        vectors = []
        self.id_map = []
        corpus = []
        for article in articles:
            vector = self.embedder.encode(article.content)
            vectors.append(vector)
            self.id_map.append(article.id)
            corpus.append(article.content.split())

        if not articles and use_fallback_dataset:
            dataset = load_dataset("pubmed", split="train[:1000]")
            for doc in dataset:
                content = doc["abstract"] + "\n" + doc["text"]
                vector = self.embedder.encode(content)
                vectors.append(vector)
                self.id_map.append(uuid4())
                corpus.append(content.split())

        if vectors:
            vectors = np.array(vectors).astype("float32")
            self.index = faiss.IndexFlatL2(self.dimension)
            faiss.normalize_L2(vectors)
            self.index.add(vectors)
            self.bm25 = BM25Okapi(corpus)

    def retrieve_hybrid(self, query: str, top_k: int = 5) -> list:
        if self.index is None:
            return []

        query_vector = self.embedder.encode(query)
        query_vector = np.array([query_vector]).astype("float32")
        faiss.normalize_L2(query_vector)
        distances_sem, indices_sem = self.index.search(query_vector, top_k)

        tokenized_query = query.split()
        doc_scores = self.bm25.get_scores(tokenized_query)
        indices_key = np.argsort(doc_scores)[-top_k:]

        hybrid_indices = list(set(list(indices_sem[0]) + list(indices_key)))
        retrieved_ids = [
            self.id_map[idx] for idx in hybrid_indices if idx < len(self.id_map)
        ]
        return retrieved_ids[:top_k]

    async def fetch_external_knowledge(self, query: str, top_k: int = 3) -> list:

        try:
            response = requests.get(
                f"https://api.crossref.org/works?query={query}&rows={top_k}", timeout=5
            )
            if response.status_code == 200:
                items = response.json().get("message", {}).get("items", [])
                return [
                    item.get("abstract", "") or item.get("title", "") for item in items
                ]
        except Exception as e:
            logger.error(f"Error fetching external knowledge: {e}")
            return []
        return []

    def add_memory(self, context: str):
        self.memory.append(context)


class Agent:
    def __init__(self, model, tokenizer):
        self.model = model
        self.tokenizer = tokenizer
        self.cache = {}

    def act(self, prompt: str, max_len: int = 1500) -> str:

        cache_key = hash(prompt)
        if cache_key in self.cache:
            return self.cache[cache_key]

        inputs = self.tokenizer(prompt, return_tensors="pt")
        outputs = self.model.generate(
            **inputs,
            max_length=max_len,
            num_beams=5,
            temperature=0.7,
            use_cache=True,
        )
        result = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
        self.cache[cache_key] = result
        return result


class PlannerAgent(Agent):
    def __init__(self) -> None:
        quantization_config = BitsAndBytesConfig(load_in_4bit=True)
        model = AutoModelForCausalLM.from_pretrained(
            "meta-llama/Llama-3-70B", quantization_config=quantization_config
        )
        tokenizer = AutoTokenizer.from_pretrained("meta-llama/Llama-3-70B")
        lora_config = LoraConfig(
            r=16, lora_alpha=32, target_modules=["q_proj", "v_proj", "k_proj", "o_proj"]
        )
        model = prepare_model_for_kbit_training(model)
        model = get_peft_model(model, lora_config)
        super().__init__(model, tokenizer)

    def generate_outline(self, text: str) -> list:
        prompt = f"Act as planner agent: Generate detailed outline for scientific article based on: {text}. Include abstract, keywords, introduction, methods, results, discussion, conclusion, references, and figure descriptions."
        outline_text = self.act(prompt, max_len=400)
        sections = [s.strip() for s in outline_text.split("\n") if s.strip()]
        return sections


class GeneratorAgent(Agent):
    def __init__(self) -> None:
        quantization_config = BitsAndBytesConfig(load_in_4bit=True)
        model = AutoModelForCausalLM.from_pretrained(
            "meta-llama/Llama-3-70B", quantization_config=quantization_config
        )
        tokenizer = AutoTokenizer.from_pretrained("meta-llama/Llama-3-70B")
        lora_config = LoraConfig(
            r=16, lora_alpha=32, target_modules=["q_proj", "v_proj", "k_proj", "o_proj"]
        )
        model = prepare_model_for_kbit_training(model)
        model = get_peft_model(model, lora_config)
        super().__init__(model, tokenizer)

    def generate_section(
        self,
        section: str,
        context: str,
        header: str,
        control: str = "medium",
        user_degree: str = "bachelor",
    ) -> str:
        complexity = (
            "advanced" if user_degree.lower() in ["phd", "master"] else "simple"
        )
        prompt = f"Act as generator agent: Generate detailed {control}-length section '{section}' for article '{header}'. Use context: {context}. Ensure scientific accuracy, {complexity} terminology, and add figure descriptions and LaTeX equations if needed."
        text = self.act(prompt)
        if "results" in section.lower():
            text += "\n" + self.generate_visualization(section)
        return text

    def generate_visualization(self, section: str) -> str:
        fig, ax = plt.subplots()
        x = np.linspace(0, 10, 100)
        y = np.sin(x)
        ax.plot(x, y, label="Sample Data")
        ax.set_title(f"Visualization for {section}")
        ax.set_xlabel("X-axis")
        ax.set_ylabel("Y-axis")
        ax.legend()
        buffer = BytesIO()
        plt.savefig(buffer, format="png")
        plt.close()
        buffer.seek(0)
        file_path = save_file(f"figure_{uuid4()}.png", buffer)
        return f"Figure: Sample plot saved at {file_path}"


class CriticAgent:
    def __init__(self) -> None:
        self.classifier = DebertaV2ForSequenceClassification.from_pretrained(
            "microsoft/deberta-v3-large"
        )
        self.tokenizer = DebertaV2Tokenizer.from_pretrained(
            "microsoft/deberta-v3-large"
        )
        self.coherence_model = CoherenceModel()
        self.embedder = SentenceTransformer("all-mpnet-base-v2")
        self.rouge = load("rouge")
        self.bleu = load("bleu")

    def evaluate_coherence(self, text: str) -> float:
        sentences = sent_tokenize(text)
        embeddings = self.embedder.encode(sentences)
        embeddings_tensor = torch.tensor(embeddings).float()
        scores = self.coherence_model(embeddings_tensor).mean().item()
        return scores

    def critique(self, text: str, reference: str = None) -> dict:
        inputs = self.tokenizer(text[:512], return_tensors="pt")
        outputs = self.classifier(**inputs)
        label = outputs.logits.argmax().item()
        coherence = self.evaluate_coherence(text)
        metrics = {}
        if reference:
            metrics["rouge"] = self.rouge.compute(
                predictions=[text], references=[reference]
            )
            metrics["bleu"] = self.bleu.compute(
                predictions=[text], references=[[reference]]
            )
        critique = (
            "Needs improvement: low coherence."
            if label == 0 or coherence < 0.8
            else "Good quality."
        )
        return {"critique": critique, "coherence": coherence, "metrics": metrics}


class Retriever:
    def __init__(self, knowledge_base: KnowledgeBase, db_session: AsyncSession) -> None:
        self.knowledge_base = knowledge_base
        self.db_session = db_session

    async def build_context(self, section: str, iterations=3) -> str:
        context = ""
        query = section
        for _ in range(iterations):
            retrieved_ids = self.knowledge_base.retrieve_hybrid(query)
            async with self.db_session as session:
                result = await session.execute(
                    select(ArticleModel).where(ArticleModel.id.in_(retrieved_ids))
                )
                articles = result.scalars().all()
            new_context = "\n".join([article.content for article in articles])
            external_context = await self.knowledge_base.fetch_external_knowledge(query)
            context += new_context + "\n" + "\n".join(external_context) + "\n"
            query = f"Refine query using chain-of-thought: {query} based on new info {new_context[:200]}"
            self.knowledge_base.add_memory(context)
        return context


class PostProcessor:
    def __init__(self) -> None:
        self.critic = CriticAgent()

    def process(self, text: str) -> str:
        sentences = sent_tokenize(text)
        unique_sentences = []
        for sent in sentences:
            if sent not in unique_sentences:
                unique_sentences.append(sent)
        text = " ".join(unique_sentences).strip()
        score = self.critic.evaluate_coherence(text)
        if score < 0.8:
            text = f"[Refined for coherence: {score}] " + text
        return text

    def check_plagiarism(self, text: str, reference_texts: list) -> float:
        max_similarity = 0
        for ref in reference_texts:
            seq = difflib.SequenceMatcher(None, text, ref)
            similarity = seq.ratio()
            max_similarity = max(max_similarity, similarity)
        return max_similarity

    def generate_references(self, query: str) -> str:
        external_refs = KnowledgeBase(None).fetch_external_knowledge(query, top_k=5)
        refs = [f"{i+1}. {ref[:200]}..." for i, ref in enumerate(external_refs)]
        return "References:\n" + "\n".join(refs)

    def generate_figure_desc(self, section: str) -> str:
        return f"Figure for {section}: A diagram illustrating key concepts, generated using Matplotlib."


class ArticleOperation:
    def __init__(self, db_session: AsyncSession) -> None:
        self.db_session = db_session
        self.knowledge_base = KnowledgeBase(db_session)
        self.planner_agent = PlannerAgent()
        self.generator_agent = GeneratorAgent()
        self.critic_agent = CriticAgent()
        self.retriever = Retriever(self.knowledge_base, db_session)
        self.post_processor = PostProcessor()

    async def fine_tune_lora(self):
        dataset = load_dataset("pubmed", split="train[:1000]")
        model = self.generator_agent.model
        tokenizer = self.generator_agent.tokenizer
        trainer = SFTTrainer(
            model=model,
            tokenizer=tokenizer,
            train_dataset=dataset,
            dataset_text_field="abstract",
            max_seq_length=512,
        )
        trainer.train()
        logger.info("LoRA fine-tuned with PubMed dataset.")

    async def generate_article(self, user_id: UUID, text_id: UUID, header: str) -> dict:
        await self.fine_tune_lora()

        async with self.db_session as session:
            user = await session.execute(
                select(UserModel).where(UserModel.id == user_id)
            )
            user = user.scalars().first()
            if not user:
                raise HTTPException(status_code=404, detail="User not found")

            text = await session.execute(
                select(TextModel).where(TextModel.id == text_id)
            )
            text = text.scalars().first()
            if not text:
                raise HTTPException(status_code=404, detail="Text not found")

        await self.knowledge_base.build_index(use_fallback_dataset=True)
        outline = self.planner_agent.generate_outline(text.text)

        generated_sections = {}
        for section in outline:
            context = await self.retriever.build_context(section)
            section_text = self.generator_agent.generate_section(
                section,
                context,
                header,
                control=user.degree.lower(),
                user_degree=user.degree,
            )
            critique = self.critic_agent.critique(section_text)
            if "improvement" in critique["critique"]:
                section_text = self.generator_agent.generate_section(
                    section,
                    context + critique["critique"],
                    header,
                    user_degree=user.degree,
                )
            generated_sections[section.lower()] = (
                section_text + "\n" + self.post_processor.generate_figure_desc(section)
            )

        generated_text = f"Title: {header}\n\n"
        generated_text += "Abstract: " + generated_sections.get("abstract", "") + "\n\n"
        generated_text += (
            "Keywords: "
            + ", ".join(generated_sections.get("keywords", "").split())
            + "\n\n"
        )
        for sec in ["introduction", "methods", "results", "discussion", "conclusion"]:
            if sec in generated_sections:
                generated_text += f"{sec.capitalize()}: {generated_sections[sec]}\n\n"
        generated_text += await self.post_processor.generate_references(header)
        generated_text = self.post_processor.process(generated_text)

        async with self.db_session as session:
            new_article = AIModel(
                title=f"Generated Article: {header}",
                model_output=generated_text,
                user_id=user.id,
                text_id=text.id,
                article_id=uuid4(),
                created_at=datetime.utcnow(),
                translated_output=None,
            )
            session.add(new_article)
            await session.commit()

        return {
            "success": "Article generated successfully",
            "article_id": str(new_article.id),
        }

    async def refine_article(self, article_id: UUID, user_id: UUID) -> dict:
        async with self.db_session as session:
            article = await session.execute(
                select(AIModel).where(
                    AIModel.id == article_id, AIModel.user_id == user_id
                )
            )
            article = article.scalars().first()
            if not article:
                raise HTTPException(status_code=404, detail="Article not found")

        refined_text = article.model_output
        for _ in range(3):
            critique = self.critic_agent.critique(refined_text)
            if critique["coherence"] >= 0.8:
                break
            refined_text = self.generator_agent.generate_section(
                "full article",
                refined_text + "\nCritique: " + critique["critique"],
                article.title,
                control="detailed",
            )
            refined_text = self.post_processor.process(refined_text)

        article.model_output = refined_text
        await self.db_session.commit()

        return {
            "success": "Article refined successfully",
            "article_id": str(article.id),
        }

    async def evaluate_article(self, article_id: UUID, user_id: UUID) -> dict:
        async with self.db_session as session:
            article = await session.execute(
                select(AIModel).where(
                    AIModel.id == article_id, AIModel.user_id == user_id
                )
            )
            article = article.scalars().first()
            if not article:
                raise HTTPException(status_code=404, detail="Article not found")

        coherence_score = self.critic_agent.evaluate_coherence(article.model_output)
        critique = self.critic_agent.critique(article.model_output)
        return {
            "coherence_score": coherence_score,
            "critique": critique["critique"],
            "metrics": critique["metrics"],
        }

    async def delete_article(self, article_id: UUID) -> None:
        async with self.db_session as session:
            article = await session.execute(
                select(AIModel).where(AIModel.id == article_id)
            )
            article = article.scalars().first()
            if not article:
                raise HTTPException(status_code=404, detail="Article not found.")
            await session.delete(article)
            await session.commit()


class FileOperation:
    def __init__(self, db_session: AsyncSession) -> None:
        self.db_session = db_session

    async def generate_pdf(self, article_id: UUID, user_id: UUID) -> dict:
        async with self.db_session as session:
            article = await session.execute(
                select(AIModel).where(
                    AIModel.id == article_id, AIModel.user_id == user_id
                )
            )
            article = article.scalars().first()
            if not article:
                raise HTTPException(
                    status_code=404, detail="Article not found or access denied."
                )

            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(190, 10, article.model_output)

            pdf_buffer = BytesIO()
            pdf_output = pdf.output(dest="S").encode("latin-1")
            pdf_buffer.write(pdf_output)
            pdf_buffer.seek(0)

            file_name = f"article_{article.id}.pdf"
            file_path = save_file(file_name, pdf_buffer)

            new_file = FileModel(
                ai_model_id=article.id,
                user_id=user_id,
                file_name=file_name,
                file_path=file_path,
                file_type="pdf",
                created_at=datetime.utcnow(),
            )
            session.add(new_file)
            await session.commit()

            return {"download_url": f"/download/{new_file.id}"}

    async def generate_word(self, article_id: UUID, user_id: UUID) -> dict:
        async with self.db_session as session:
            article = await session.execute(
                select(AIModel).where(
                    AIModel.id == article_id, AIModel.user_id == user_id
                )
            )
            article = article.scalars().first()
            if not article:
                raise HTTPException(
                    status_code=404, detail="Article not found or access denied."
                )

            doc = Document()
            doc.add_heading(article.title, level=1)
            doc.add_paragraph(article.model_output)

            word_buffer = BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            file_name = f"article_{article.id}.docx"
            file_path = save_file(file_name, word_buffer)

            new_file = FileModel(
                ai_model_id=article.id,
                user_id=user_id,
                file_name=file_name,
                file_path=file_path,
                file_type="word",
                created_at=datetime.utcnow(),
            )
            session.add(new_file)
            await session.commit()

            return {"download_url": f"/download/{new_file.id}"}

    async def generate_latex(self, article_id: UUID, user_id: UUID) -> dict:
        async with self.db_session as session:
            article = await session.execute(
                select(AIModel).where(
                    AIModel.id == article_id, AIModel.user_id == user_id
                )
            )
            article = article.scalars().first()
            if not article:
                raise HTTPException(
                    status_code=404, detail="Article not found or access denied."
                )

            latex_content = latexencode(article.model_output)
            buffer = BytesIO()
            buffer.write(latex_content.encode("utf-8"))
            buffer.seek(0)

            file_name = f"article_{article.id}.tex"
            file_path = save_file(file_name, buffer)

            new_file = FileModel(
                ai_model_id=article.id,
                user_id=user_id,
                file_name=file_name,
                file_path=file_path,
                file_type="latex",
                created_at=datetime.utcnow(),
            )
            session.add(new_file)
            await session.commit()

            return {"download_url": f"/download/{new_file.id}"}

    async def get_file(self, file_id: UUID, user_id: UUID) -> dict:
        async with self.db_session as session:
            file = await session.execute(
                select(FileModel).where(
                    FileModel.id == file_id, FileModel.user_id == user_id
                )
            )
            file = file.scalars().first()
            if not file:
                raise HTTPException(
                    status_code=404, detail="File not found or access denied."
                )

            return {"file_name": file.file_name, "file_path": file.file_path}


class TranslationOperation:
    def __init__(self, db_session: AsyncSession):
        self.db_session = db_session

    async def translate_article(self, article_id: UUID):
        result = await self.db_session.execute(
            select(AIModel).where(AIModel.id == article_id)
        )
        ai_article = result.scalar_one_or_none()

        if not ai_article:
            raise HTTPException(status_code=404, detail="Article not found.")

        if not ai_article.model_output:
            raise HTTPException(status_code=400, detail="No content to translate.")

        original_text = ai_article.model_output
        chunk_size = 3000
        chunks = [
            original_text[i : i + chunk_size]
            for i in range(0, len(original_text), chunk_size)
        ]

        translated_chunks = []
        for chunk in chunks:
            translated = GoogleTranslator(source="auto", target="fa").translate(chunk)
            translated_chunks.append(translated)

        full_translation = "\n".join(translated_chunks)
        ai_article.translated_output = full_translation
        await self.db_session.commit()

        return {
            "detail": "Translated successfully.",
            "translated_text": full_translation,
        }
